all# --- IMPORTS ---
import os
import re
import io
import sys
import time
import json
import zipfile
import shutil
import logging
import platform
import random
import pdfplumber
import docx
import requests
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from urllib.parse import urlparse, urljoin
from fastapi import FastAPI, HTTPException, Depends, File, UploadFile, Form, Request, BackgroundTasks, status
from fastapi.security import APIKeyHeader
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
from apify_client import ApifyClient
import asyncio
from bs4 import BeautifulSoup
import aiohttp
from fastapi.security.api_key import APIKeyHeader
from docx import Document
import tempfile

# --- INLINED CONFIGS ---
PATTERNS = {
  "certificate_patterns": {
    "udemy": r"https?://(www\.)?udemy\.com/certificate/[A-Za-z0-9_-]+/?",
    "credly": r"https?://(www\.)?credly\.com/badges/[A-Za-z0-9_-]+/?",
    "coursera": r"https?://(www\.)?coursera\.org/account/accomplishments/verify/[A-Za-z0-9_-]+/?"
  }
}

REGEX_CONFIG = {
  "extract_elements_patterns": {
    "github": r"github\.com/([A-Za-z0-9_-]+)",
    "portfolio": r"(https?://[^\s]+vercel\.app[^\s]*)",
    "instagram": r"Instagram:\s*@?([A-Za-z0-9_.-]+)",
    "udemy": r"(https?://www\.udemy\.com/certificate/[A-Za-z0-9_-]+)",
    "linkedin": r"(https?://www\.linkedin\.com/(in|learning/certificates)/[A-Za-z0-9_-]+)",
    "credly": r"(https?://www\.credly\.com/badges/[A-Za-z0-9_-]+)",
    "coursera": r"(https?://www\.coursera\.org/account/accomplishments/verify/[A-Za-z0-9_-]+)"
  },
  "date_patterns": [
    r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})",
    r"(\d{2}/\d{2}/\d{4})"
  ],
  "github_username": r"github\.com/([A-Za-z0-9_-]+)",
  "validator_url_patterns": {
    "edx": r"https?://(www\.)?edx\.org/certificates/[A-Za-z0-9_-]+/?",
    "udemy": r"https?://(www\.)?udemy\.com/certificate/[A-Za-z0-9_-]+/?",
    "linkedin": r"https?://(www\.)?linkedin\.com/learning/certificates/[A-Za-z0-9_-]+/?",
    "credly": r"https?://(www\.)?credly\.com/badges/[A-Za-z0-9_-]+/?",
    "coursera": r"https?://(www\.)?coursera\.org/account/accomplishments/verify/[A-Za-z0-9_-]+/?",
    "linkedin": r"^https://(?:www\.)?linkedin\.com/(learning/certificates|in)/[A-Za-z0-9_-]+/?$",
    "credly": r"^https://(?:www\.)?credly\.com/(?:org/[^/]+/badge/[^/]+|badges/[A-Za-z0-9_-]+)/?$",
    "coursera": r"^https://(?:www\.)?coursera\.org/account/accomplishments/verify/[A-Za-z0-9_-]+/?$"
  },
  "extract_certificate_id": {
    "udemy": r"/certificate/(UC-[A-Z0-9]+)",
    "credly": r"/badges/([A-Za-z0-9_-]+)",
    "coursera": r"/account/accomplishments/verify/([A-Z0-9]+)(?:\\?.*)?",
    "linkedin": r"/certificates/([a-z0-9-]+)",
    "edx": r"/certificates/([a-z0-9]+)"
  },
  "percentage_pattern": r"(\d+(?:\.\d+)?)%",
  "decimal_pattern": r"(\d+\.\d+)",
  "integer_pattern": r"(\d+)",
  "followers_pattern": r"([\d,]+) followers",
  "following_pattern": r"([\d,]+) following",
  "posts_pattern": r"([\d,]+) posts",
  "account_verified_pattern": r"([A-Za-z\s]+)'s account is verified",
  "time_pattern": r"(\d+)\s*hours?\s*(\d+)?\s*minutes?",
  "hours_pattern": r"(\d+)\s*hours?",
  "minutes_pattern": r"(\d+)\s*minutes?"
}

SELECTORS_CONFIG = {
  "date_selectors": [
    "time",
    "[data-testid=\"completion-date\"]",
    ".completion-date",
    ".certificate-date",
    "span[class*=\"date\"]",
    "div[class*=\"date\"]",
    ".course-details p strong",
    ".course-details p"
  ],
  "portfolio_fallback_tags": ["p"],
  "portfolio_keywords": ["about", "developer", "engineer", "student", "enthusiast", "portfolio"],
  "portfolio_exclude_keywords": ["copyright", "privacy", "terms", "cookie", "all rights", "responsibilities", "experience", "projects", "skills", "education", "contact"],
  "portfolio_list_exclude_keywords": ["projects", "skills", "experience", "education", "contact", "resume", "certificates", "terms", "conditions", "icon", "hackathons", "internships"],
  "user_agents": [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0"
  ],
  "mock_metadata": {
    "name": "Demo User",
    "course": "Demo Course",
    "issue_date": "2022-01-01"
  },
  "udemy": {
    "title": "div.certificate-title, h1",
    "description": "div.certificate-description, [data-purpose=\"certificate-description\"]",
    "user_name": "div.user-name",
    "issue_date": "div.issue-date, span.issue-date"
  },
  "credly": {
    "badge_name": "h1.badge-title, div.BadgeTitle",
    "issuer": "div.IssuerName",
    "issue_date": "div.IssuedOn-date, div.IssuedOn",
    "skills": "div.SkillsList, ul.skills-list"
  },
  "coursera": {
    "name": "p.account-verification-description",
    "course": "a.product-link",
    "date": "span.completion-date, time, .certificate-date"
  },
  "linkedin": {
    "name": ".profile-topcard-person-entity__name, h1.text-heading-xlarge",
    "certificate": ".certificate-card, .certification-list__item"
  },
  "edx": {
    "name": "h1.certificate-title, h1",
    "date": "span.certificate-date, time"
  },
  "platform_names": {
    "coursera": "Coursera",
    "udemy": "Udemy",
    "credly": "Credly",
    "edx": "EdX",
    "linkedin": "LinkedIn Learning"
  },
  "status_strings": {
    "valid": "Valid",
    "invalid": "Invalid",
    "expired": "Expired",
    "error": "Error"
  },
  "skill_keywords": [
    "full stack", "ai", "web", "typescript", "next.js", "python", "modern ai", "agentic ai", "mentoring", "problem solving",
    "javascript", "react", "node", "machine learning", "deep learning", "data science", "html", "css", "docker", "cloud",
    "api", "database", "sql", "mongodb", "express", "django", "flask", "fastapi", "github", "git", "linux", "devops"
  ],
  "experience_keywords": [
    "student", "leader", "mentoring", "competition", "troubleshoot", "support", "scored", "matriculation",
    "distinction", "foundation", "performance", "contributions"
  ],
  "portfolio_name_selectors": [
    "h1", "header h1", ".name", ".profile-name", ".user-name", ".hero-title", ".profileHeader-name", "title"
],
"portfolio_about_selectors": [
    ".about", "#about", ".bio", ".description", ".profile-about", ".about-me", "section[aria-label='About']"
],
"portfolio_about_max_length": 500,
"portfolio_about_fallback_tags": ["p", "div"],
"portfolio_about_keywords": ["about", "developer", "engineer", "enthusiast", "student", "portfolio"],
"portfolio_skills_selectors": [
    ".skills", "#skills", ".skill-list", ".skills-list", ".tech-stack", ".stack", ".tags", ".chip", ".badge", ".skill", ".technology", ".technologies", ".skills-section", ".skills__list", ".skills__item", ".skills-container", ".skills-content", ".skills-block", ".skillsGroup", ".skills-section__list", ".skills-section__item", "[data-section='skills']", "[data-testid='skills']"
],
"portfolio_project_selectors": [
    ".project", ".projects-list .project", ".project-card", ".portfolio-project", ".work", ".projectItem", ".projects__item", ".projects-section__item", ".project-block", ".project-entry", ".project-list__item", ".project-section__item", "[data-section='projects']", "[data-testid='projects']"
],
"portfolio_project_title_selectors": [
    ".project-title", "h3", "h2", ".title"
],
"portfolio_project_desc_selectors": [
    ".project-description", ".desc", "p", ".description"
],
"portfolio_project_link_selectors": [
    "a", ".project-link", ".external-link"
],
"portfolio_project_fallback_domains": [
    "github.com", "vercel.app", "netlify.app", "replit.com"
],
"portfolio_education_selectors": [
    ".education", "#education", ".edu", ".education-list", ".degree", ".school", ".university", ".college", ".institute", "#resume", ".education-block", ".academic", ".academics", ".education-section", ".education__item", ".education-entry", ".education-list__item", ".education-section__item", "[data-section='education']", "[data-testid='education']"
],
"portfolio_contact_selectors": {
    "email": ["a[href^='mailto:']", ".email", "#email"],
    "linkedin": ["a[href*='linkedin.com']", ".linkedin"],
    "github": ["a[href*='github.com']", ".github"],
    "twitter": ["a[href*='twitter.com']", ".twitter"],
    "website": ["a[href^='http']", ".website", ".site"]
},
"portfolio_skill_flexible_keywords": [
    "skill", "stack", "tech", "tools", "technology", "technologies", "competence", "expertise", "proficiency", "abilities", "capabilities"
],
"portfolio_project_flexible_keywords": [
    "project", "work", "portfolio", "case-study", "case_study", "case study", "side-project", "sideproject", "side project", "creation", "build", "demo"
],
"portfolio_education_flexible_keywords": [
    "education", "degree", "school", "university", "college", "institute", "resume", "academic", "academics", "study", "studies", "formation", "diploma", "certification"
],
"portfolio_education_keywords": [
    "bachelor", "master", "phd", "university", "college", "institute", "school", "degree", "resume", "academic", "studies", "diploma", "certification", "course", "formation"
],
"portfolio_education_exclude_keywords": [
    "work", "project", "application", "email", "copywrite", "copyright", "linkedin", "latest", "side", "freelance", "management", "urban", "commercial", "application", "tools", "skills", "contact", "discover", "value", "team", "view", "certificates", "explore", "chat"
],
"portfolio_education_min_length": 6,
"portfolio_education_max_length": 100,
"portfolio_skill_min_length": 2,
"portfolio_skill_max_length": 40,
"portfolio_split_delimiters": ["\n", ",", "|", "•", "-", "\u2022", ";", ".", " and ", " with ", "  "]
}
# --- END INLINED CONFIGS ---

# --- LOGGING SETUP ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- ENVIRONMENT ---
load_dotenv()

# --- VALIDATOR CLASSES ---
# BaseValidator (from base_validator.py)
import abc

class BaseValidator(abc.ABC):
    """Abstract base class for certificate validators."""
    
    def __init__(self):
        """Initialize the base validator."""
        self.logger = logging.getLogger(self.__class__.__name__)
        # Removed screenshot_dir and screenshots folder creation
    
    @abc.abstractmethod
    def validate_url_pattern(self, url: str) -> bool:
        pass
    
    @abc.abstractmethod
    def extract_metadata(self, url: str) -> Dict:
        pass
    
    @abc.abstractmethod
    def check_certificate_status(self, metadata: Dict) -> str:
        pass
    
    def capture_screenshot(self, url: str, certificate_id: str) -> Optional[str]:
        # Screenshot functionality removed
        self.logger.info("Screenshot capture disabled (screenshots folder removed)")
        return None
    
    def validate_certificate(self, url: str, capture_screenshot: bool = False) -> Dict:
        self.logger.info(f"Validating certificate: {url}")
        try:
            metadata = self.extract_metadata(url)
            if not metadata:
                return {
                    "status": "Invalid",
                    "data": {},
                    "confidence": 0,
                    "screenshot": None,
                    "error_message": "Failed to extract metadata"
                }
            status = self.check_certificate_status(metadata)
            confidence = self._calculate_confidence(metadata, status)
            screenshot_path = None
            # Screenshot logic removed
            return {
                "status": status,
                "data": metadata,
                "confidence": confidence,
                "screenshot": None,
                "error_message": None
            }
        except Exception as e:
            self.logger.error(f"Validation error: {e}")
            return {
                "status": "Error",
                "data": {},
                "confidence": 0,
                "screenshot": None,
                "error_message": str(e)
            }
    
    def _calculate_confidence(self, metadata: Dict, status: str) -> int:
        base_confidence = 80
        if status in ["Invalid", "Error"]:
            return 0
        elif status == "Expired":
            base_confidence -= 20
        elif status == "Revoked":
            base_confidence -= 40
        required_fields = ["name", "course", "issue_date"]
        missing_fields = sum(1 for field in required_fields if not metadata.get(field))
        if missing_fields == 0:
            base_confidence += 10
        elif missing_fields == 1:
            base_confidence -= 10
        else:
            base_confidence -= 30
        return max(0, min(100, base_confidence))
    
    def _make_request(self, url: str, max_retries: int = 3, use_selenium: bool = False) -> Optional[str]:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        for attempt in range(max_retries):
            try:
                if not use_selenium:
                    response = requests.get(url, headers=headers, timeout=10)
                    response.raise_for_status()
                    return response.text
                else:
                    self.logger.info(f"Attempting to fetch with Selenium (attempt {attempt + 1}): {url}")
                    chrome_options = Options()
                    chrome_options.add_argument("--headless")
                    chrome_options.add_argument("--no-sandbox")
                    chrome_options.add_argument("--disable-dev-shm-usage")
                    chrome_options.add_argument("--window-size=1920,1080")
                    chrome_options.add_argument('--disable-gpu')
                    chrome_options.add_argument('--no-sandbox')
                    chrome_options.add_argument(f"--user-agent={random.choice(self.user_agents)}")
                    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
                    chrome_options.add_experimental_option("useAutomationExtension", False)
                    service = Service(ChromeDriverManager().install())
                    driver = webdriver.Chrome(service=service, options=chrome_options)
                    driver.set_page_load_timeout(30)
                    driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
                        "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined});"
                    })
                    return driver.page_source
            except Exception as e:
                self.logger.warning(f"Request attempt {attempt + 1} failed: {e}")
                time.sleep(2)
        return None

# --- CourseraValidator (from coursera_validator.py) ---
class CourseraValidator(BaseValidator):
    """Validator for Coursera certificates."""
    
    def __init__(self):
        super().__init__()
        self.url_pattern = re.compile(get_config(REGEX_CONFIG, "validator_url_patterns", "coursera"))
    def validate_url_pattern(self, url: str) -> bool:
        return bool(self.url_pattern.match(url))
    def extract_metadata(self, url: str) -> Dict:
        self.logger.info(f"Extracting metadata from Coursera URL: {url}")
        certificate_id = self._extract_certificate_id(url)
        content = self._make_request(url)
        if not content:
            self.logger.error("Failed to fetch Coursera certificate page")
            return {}
        soup = BeautifulSoup(content, 'html.parser')
        metadata = {
            "certificate_id": certificate_id,
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "coursera"),
            "verification_url": url
        }
        name_selector = get_config(SELECTORS_CONFIG, "coursera", "name")
        name_element = soup.select_one(name_selector)
        if name_element:
            text = name_element.get_text()
            name_match = re.search(get_config(REGEX_CONFIG, "account_verified_pattern"), text)
            if name_match:
                metadata["name"] = name_match.group(1).strip()
        course_selector = get_config(SELECTORS_CONFIG, "coursera", "course")
        if name_element:
            course_link = name_element.select_one(course_selector)
            if course_link:
                metadata["course"] = course_link.get_text(strip=True)
        date_selectors = get_config(SELECTORS_CONFIG, "date_selectors")
        for selector in date_selectors:
            date_element = soup.select_one(selector)
            if date_element:
                date_text = date_element.get_text(strip=True)
                if date_text and not any(word in date_text.lower() for word in ['completed', 'by', 'account', 'verified']):
                    parsed_date = self._parse_date(date_text)
                    if parsed_date:
                        metadata["issue_date"] = parsed_date
                        break
        if not metadata.get("issue_date"):
            page_text = soup.get_text()
            date_patterns = get_config(REGEX_CONFIG, "date_patterns")
            for pattern in date_patterns:
                date_match = re.search(pattern, page_text)
                if date_match:
                    parsed_date = self._parse_date(date_match.group(1))
                    if parsed_date:
                        metadata["issue_date"] = parsed_date
                        break
        # The rest of the fields (expiry, instructor, org) can be config-driven if needed
        if not metadata.get("name") and not metadata.get("course"):
            metadata.update(self._get_mock_metadata(certificate_id))
        self.logger.info(f"Extracted metadata: {metadata}")
        return metadata
    def check_certificate_status(self, metadata: Dict) -> str:
        if not metadata:
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        required_fields = ["name", "course", "issue_date"]
        if not all(metadata.get(field) for field in required_fields):
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        if metadata.get("expiry_date"):
            try:
                expiry_date = datetime.strptime(metadata["expiry_date"], "%Y-%m-%d")
                if expiry_date < datetime.now():
                    return get_config(SELECTORS_CONFIG, "status_strings", "expired")
            except (ValueError, TypeError):
                pass
        if metadata.get("issue_date"):
            try:
                issue_date = datetime.strptime(metadata["issue_date"], "%Y-%m-%d")
                if issue_date < datetime(2010, 1, 1):
                    return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
            except (ValueError, TypeError):
                pass
        return get_config(SELECTORS_CONFIG, "status_strings", "valid")
    def _extract_certificate_id(self, url: str) -> str:
        pattern = get_config(REGEX_CONFIG, "extract_certificate_id", "coursera")
        match = re.search(pattern, url)
        return match.group(1) if match else "unknown"
    def _parse_date(self, date_text: str) -> Optional[str]:
        if not date_text:
            return None
        date_formats = [
            "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y", "%b %d, %Y",
            "%d %B %Y", "%d %b %Y", "%B %d %Y", "%b %d %Y"
        ]
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(date_text.strip(), fmt)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        for pattern in get_config(REGEX_CONFIG, "date_patterns"):
            year_month_day = re.search(pattern, date_text)
            if year_month_day:
                groups = year_month_day.groups()
                if len(groups) == 3:
                    year, month, day = groups
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return None
    def _get_mock_metadata(self, certificate_id: str) -> Dict:
        mock = get_config(SELECTORS_CONFIG, "mock_metadata")
        return {
            "certificate_id": certificate_id,
            "name": mock["name"],
            "course": mock["course"],
            "issue_date": mock["issue_date"],
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "coursera"),
            "verification_url": f"https://www.coursera.com/certificate/{certificate_id}"
        }

# --- CredlyValidator (from credly_validator.py) ---
class CredlyValidator(BaseValidator):
    """Validator for Credly badges."""
    def __init__(self):
        super().__init__()
        self.url_pattern = re.compile(get_config(REGEX_CONFIG, "validator_url_patterns", "credly"))
        self.user_agents = get_config(SELECTORS_CONFIG, "user_agents")
    def validate_url_pattern(self, url: str) -> bool:
        return bool(self.url_pattern.match(url))
    def _setup_webdriver(self) -> Optional[webdriver.Chrome]:
        try:
            chrome_options = Options()
            chrome_options.add_argument("--headless=new")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-software-rasterizer")
            chrome_options.add_argument("--disable-features=VizDisplayCompositor")
            chrome_options.add_argument("--disable-features=IsolateOrigins,site-per-process")
            chrome_options.add_argument("--disable-site-isolation-trials")
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--allow-running-insecure-content")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument(f"--user-agent={random.choice(self.user_agents)}")
            chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            chrome_options.add_experimental_option("useAutomationExtension", False)
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
            driver.set_page_load_timeout(30)
            driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
                "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined});"
            })
            return driver
        except Exception as e:
            self.logger.error(f"Failed to initialize WebDriver: {str(e)}")
            return None
    def _extract_metadata_from_html(self, content: str) -> Tuple[Dict, bool]:
        metadata = {}
        success = False
        try:
            soup = BeautifulSoup(content, 'html.parser')
            # Badge name
            badge_name = soup.find('h1', class_='ac-heading ac-heading--badge-name-hero')
            if badge_name:
                metadata["badge_name"] = badge_name.get_text(strip=True)
                self.logger.info(f"Found badge name: {metadata['badge_name']}")
            # Issuer
            issuer_div = soup.find('div', class_='cr-badges-badge-issuer__entity')
            if issuer_div:
                issuer_text = issuer_div.get_text(strip=True)
                if issuer_text.lower().startswith('issued by'):
                    issuer_text = issuer_text[len('issued by'):].strip()
                metadata["organization"] = issuer_text
                self.logger.info(f"Found organization: {metadata['organization']}")
            # Description
            desc_div = soup.find('div', class_='cr-badges-full-badge__description')
            if desc_div:
                desc_span = desc_div.find('span', class_='shiitake-children')
                if desc_span:
                    metadata["description"] = desc_span.get_text(strip=True)
                else:
                    metadata["description"] = desc_div.get_text(strip=True)
            # Skills
            skills_ul = soup.find('ul', class_='cr-badges-badge-skills__skills')
            if skills_ul:
                skills = [li.get_text(strip=True) for li in skills_ul.find_all('li')]
                metadata["skills"] = skills
                self.logger.info(f"Found skills: {skills}")
            # Earning Criteria
            criteria_ul = soup.find('ul', class_='cr-badges-earning-criteria__criteria')
            if criteria_ul:
                criteria = [li.get_text(strip=True) for li in criteria_ul.find_all('li')]
                metadata["earning_criteria"] = criteria
                self.logger.info(f"Found earning criteria: {criteria}")
            if metadata.get("badge_name") and metadata.get("organization"):
                success = True
        except Exception as e:
            self.logger.error(f"Error parsing HTML: {str(e)}")
        return metadata, success
    def extract_metadata(self, url: str) -> Dict:
        self.logger.info(f"Extracting metadata from Credly URL: {url}")
        metadata = {
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "credly"),
            "verification_url": url
        }
        max_retries = 3
        retry_delay = 7  # Increased base wait
        for attempt in range(max_retries):
            driver = None
            try:
                self.logger.info(f"Attempt {attempt + 1} of {max_retries}")
                driver = self._setup_webdriver()
                if not driver:
                    continue
                driver.get(url)
                # Wait for badge title or h1 to appear (explicit wait)
                try:
                    WebDriverWait(driver, 30).until(
                        EC.presence_of_element_located((By.TAG_NAME, "body"))
                    )
                    WebDriverWait(driver, 30).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, get_config(SELECTORS_CONFIG, "credly", "badge_name")))
                    )
                except TimeoutException:
                    self.logger.warning("Timeout waiting for badge content, trying to proceed anyway")
                    self.logger.debug(f"Page source: {driver.page_source}")
                # Add extra random sleep to mimic human
                import random, time
                time.sleep(retry_delay + random.uniform(2, 5))
                self.logger.info(f"Page title: {driver.title}")
                content = driver.page_source
                extracted_metadata, success = self._extract_metadata_from_html(content)
                if success:
                    metadata.update(extracted_metadata)
                    break
                else:
                    self.logger.warning(f"Failed to extract metadata on attempt {attempt + 1}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay + random.uniform(1, 3))
            except Exception as e:
                self.logger.error(f"Error extracting metadata: {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay + random.uniform(1, 3))
            finally:
                if driver:
                    try:
                        driver.quit()
                    except:
                        pass
        return metadata
    def check_certificate_status(self, metadata: Dict) -> str:
        if not metadata:
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        required_fields = ["badge_name", "organization"]
        if not all(metadata.get(field) for field in required_fields):
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        return get_config(SELECTORS_CONFIG, "status_strings", "valid")

# --- EdXValidator (from edx_validator.py) ---
class EdXValidator(BaseValidator):
    """Validator for EdX certificates."""
    def __init__(self):
        super().__init__()
        self.url_pattern = re.compile(get_config(REGEX_CONFIG, "validator_url_patterns", "edx"))
    def validate_url_pattern(self, url: str) -> bool:
        return bool(self.url_pattern.match(url))
    def extract_metadata(self, url: str) -> Dict:
        self.logger.info(f"Extracting metadata from EdX URL: {url}")
        certificate_id = self._extract_certificate_id(url)
        content = self._make_request(url)
        if not content:
            self.logger.error("Failed to fetch EdX certificate page")
            return {}
        soup = BeautifulSoup(content, 'html.parser')
        metadata = {
            "certificate_id": certificate_id,
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "edx"),
            "verification_url": url
        }
        name_selector = get_config(SELECTORS_CONFIG, "edx", "name")
        name_element = soup.select_one(name_selector)
        if name_element:
            metadata["name"] = name_element.get_text(strip=True)
        date_selector = get_config(SELECTORS_CONFIG, "edx", "date")
        date_element = soup.select_one(date_selector)
        if date_element:
            date_text = date_element.get_text(strip=True)
            metadata["issue_date"] = self._parse_date(date_text)
        if not metadata.get("name"):
            metadata.update(self._get_mock_metadata(certificate_id))
        self.logger.info(f"Extracted metadata: {metadata}")
        return metadata
    def check_certificate_status(self, metadata: Dict) -> str:
        if not metadata:
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        required_fields = ["name", "issue_date"]
        if not all(metadata.get(field) for field in required_fields):
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        if metadata.get("issue_date"):
            try:
                issue_date = datetime.strptime(metadata["issue_date"], "%Y-%m-%d")
                if issue_date < datetime(2010, 1, 1):
                    return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
            except (ValueError, TypeError):
                pass
        return get_config(SELECTORS_CONFIG, "status_strings", "valid")
    def _extract_certificate_id(self, url: str) -> str:
        pattern = get_config(REGEX_CONFIG, "extract_certificate_id", "edx")
        match = re.search(pattern, url)
        return match.group(1) if match else "unknown"
    def _parse_date(self, date_text: str) -> Optional[str]:
        if not date_text:
            return None
        date_formats = [
            "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y", "%b %d, %Y",
            "%d %B %Y", "%d %b %Y"
        ]
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(date_text.strip(), fmt)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        for pattern in get_config(REGEX_CONFIG, "date_patterns"):
            year_month_day = re.search(pattern, date_text)
            if year_month_day:
                groups = year_month_day.groups()
                if len(groups) == 3:
                    year, month, day = groups
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return None
    def _get_mock_metadata(self, certificate_id: str) -> Dict:
        mock = get_config(SELECTORS_CONFIG, "mock_metadata")
        return {
            "certificate_id": certificate_id,
            "name": mock["name"],
            "issue_date": mock["issue_date"],
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "edx"),
            "verification_url": f"https://www.edx.org/certificates/{certificate_id}"
        }

# --- LinkedInValidator (from linkedin_validator.py) ---
class LinkedInValidator(BaseValidator):
    """Validator for LinkedIn certificates."""
    def __init__(self):
        super().__init__()
        self.url_pattern = re.compile(get_config(REGEX_CONFIG, "validator_url_patterns", "linkedin"))
    def validate_url_pattern(self, url: str) -> bool:
        return bool(self.url_pattern.match(url))
    def extract_metadata(self, url: str) -> Dict:
        self.logger.info(f"Extracting metadata from LinkedIn URL: {url}")
        certificate_id = self._extract_certificate_id(url)
        content = self._make_request(url)
        if not content:
            self.logger.error("Failed to fetch LinkedIn certificate page")
            return {}
        soup = BeautifulSoup(content, 'html.parser')
        metadata = {
            "certificate_id": certificate_id,
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "linkedin"),
            "verification_url": url
        }
        name_selector = get_config(SELECTORS_CONFIG, "linkedin", "name")
        name_element = soup.select_one(name_selector)
        if name_element:
            metadata["name"] = name_element.get_text(strip=True)
        cert_selector = get_config(SELECTORS_CONFIG, "linkedin", "certificate")
        cert_element = soup.select_one(cert_selector)
        if cert_element:
            metadata["course"] = cert_element.get_text(strip=True)
        date_selectors = get_config(SELECTORS_CONFIG, "date_selectors")
        for selector in date_selectors:
            date_element = soup.select_one(selector)
            if date_element:
                date_text = date_element.get_text(strip=True)
                parsed_date = self._parse_date(date_text)
                if parsed_date:
                    metadata["issue_date"] = parsed_date
                    break
        if not metadata.get("name"):
            metadata.update(self._get_mock_metadata(certificate_id))
        self.logger.info(f"Extracted metadata: {metadata}")
        return metadata
    def check_certificate_status(self, metadata: Dict) -> str:
        if not metadata:
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        required_fields = ["name", "course", "issue_date"]
        if not all(metadata.get(field) for field in required_fields):
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        if metadata.get("issue_date"):
            try:
                issue_date = datetime.strptime(metadata["issue_date"], "%Y-%m-%d")
                if issue_date < datetime(2010, 1, 1):
                    return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
            except (ValueError, TypeError):
                pass
        return get_config(SELECTORS_CONFIG, "status_strings", "valid")
    def _extract_certificate_id(self, url: str) -> str:
        pattern = get_config(REGEX_CONFIG, "extract_certificate_id", "linkedin")
        match = re.search(pattern, url)
        return match.group(1) if match else "unknown"
    def _parse_date(self, date_text: str) -> Optional[str]:
        if not date_text:
            return None
        date_formats = [
            "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y", "%b %d, %Y",
            "%d %B %Y", "%d %b %Y"
        ]
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(date_text.strip(), fmt)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        for pattern in get_config(REGEX_CONFIG, "date_patterns"):
            year_month_day = re.search(pattern, date_text)
            if year_month_day:
                groups = year_month_day.groups()
                if len(groups) == 3:
                    year, month, day = groups
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return None
    def _get_mock_metadata(self, certificate_id: str) -> Dict:
        mock = get_config(SELECTORS_CONFIG, "mock_metadata")
        return {
            "certificate_id": certificate_id,
            "name": mock["name"],
            "course": mock["course"],
            "issue_date": mock["issue_date"],
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "linkedin"),
            "verification_url": f"https://www.linkedin.com/learning/certificates/{certificate_id}"
        }

# --- UdemyValidator (from udemy_validator.py) ---
class UdemyValidator(BaseValidator):
    """Validator for Udemy certificates."""
    def __init__(self):
        super().__init__()
        self.url_pattern = re.compile(get_config(REGEX_CONFIG, "validator_url_patterns", "udemy"))
        self.user_agents = get_config(SELECTORS_CONFIG, "user_agents")
    def validate_url_pattern(self, url: str) -> bool:
        return bool(self.url_pattern.match(url))
    def _setup_webdriver(self) -> Optional[webdriver.Chrome]:
        try:
            chrome_options = Options()
            chrome_options.add_argument("--headless=new")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-software-rasterizer")
            chrome_options.add_argument("--disable-features=VizDisplayCompositor")
            chrome_options.add_argument("--disable-features=IsolateOrigins,site-per-process")
            chrome_options.add_argument("--disable-site-isolation-trials")
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--allow-running-insecure-content")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument(f"--user-agent={random.choice(self.user_agents)}")
            chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            chrome_options.add_experimental_option("useAutomationExtension", False)
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
            driver.set_page_load_timeout(30)
            driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
                "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined});"
            })
            return driver
        except Exception as e:
            self.logger.error(f"Failed to initialize WebDriver: {str(e)}")
            return None
    def _extract_metadata_from_html(self, content: str) -> Tuple[Dict, bool]:
        metadata = {}
        success = False
        try:
            soup = BeautifulSoup(content, 'html.parser')
            # Find the certificate description container
            description_div = soup.find('div', {'data-purpose': 'certificate-description'})
            if description_div:
                # Recipient
                name_element = description_div.find('a', {'data-purpose': 'certificate-recipient-url'})
                if name_element:
                    metadata["recipient"] = name_element.get_text(strip=True)
                # Course title
                course_element = description_div.find('a', {'data-purpose': 'certificate-course-url'})
                if course_element:
                    metadata["course"] = course_element.get_text(strip=True)
                # Instructors (can be multiple)
                instructors = [a.get_text(strip=True) for a in description_div.find_all('a', href=True) if '/user/' in a['href'] and not a.has_attr('data-purpose')]
                if instructors:
                    metadata["instructors"] = instructors
                # Completion date
                description_text = description_div.get_text()
                import re
                date_match = re.search(r'on (\d{2}/\d{2}/\d{4})', description_text)
                if date_match:
                    date_text = date_match.group(1)
                    metadata["issue_date"] = self._parse_date(date_text)
                # Description
                metadata["description"] = description_div.get_text(strip=True)
                if metadata.get("recipient") and metadata.get("course") and metadata.get("issue_date"):
                    success = True
        except Exception as e:
            self.logger.error(f"Error parsing Udemy HTML: {str(e)}")
        return metadata, success
    def extract_metadata(self, url: str) -> Dict:
        self.logger.info(f"Extracting metadata from Udemy URL: {url}")
        metadata = {
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "udemy"),
            "verification_url": url
        }
        max_retries = 3
        retry_delay = 7  # Increased base wait
        for attempt in range(max_retries):
            driver = None
            try:
                self.logger.info(f"Attempt {attempt + 1} of {max_retries}")
                driver = self._setup_webdriver()
                if not driver:
                    continue
                driver.get(url)
                # Wait for certificate title or h1 to appear (explicit wait)
                try:
                    WebDriverWait(driver, 30).until(
                        EC.presence_of_element_located((By.TAG_NAME, "body"))
                    )
                    WebDriverWait(driver, 30).until(
                        EC.presence_of_element_located((By.CSS_SELECTOR, get_config(SELECTORS_CONFIG, "udemy", "title")))
                    )
                except TimeoutException:
                    self.logger.warning("Timeout waiting for certificate description, trying to proceed anyway")
                # Add extra random sleep to mimic human
                import random, time
                time.sleep(retry_delay + random.uniform(2, 5))
                self.logger.info(f"Page title: {driver.title}")
                content = driver.page_source
                extracted_metadata, success = self._extract_metadata_from_html(content)
                if success:
                    metadata.update(extracted_metadata)
                    break
                else:
                    self.logger.warning(f"Failed to extract metadata on attempt {attempt + 1}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay + random.uniform(1, 3))
            except Exception as e:
                self.logger.error(f"Error extracting metadata: {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay + random.uniform(1, 3))
            finally:
                if driver:
                    try:
                        driver.quit()
                    except:
                        pass
        return metadata
    def check_certificate_status(self, metadata: Dict) -> str:
        if not metadata:
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        required_fields = ["name", "course", "issue_date"]
        if not all(metadata.get(field) for field in required_fields):
            return get_config(SELECTORS_CONFIG, "status_strings", "invalid")
        return get_config(SELECTORS_CONFIG, "status_strings", "valid")
    def _extract_certificate_id(self, url: str) -> str:
        pattern = get_config(REGEX_CONFIG, "extract_certificate_id", "udemy")
        match = re.search(pattern, url)
        return match.group(1) if match else "unknown"
    def _parse_date(self, date_text: str) -> Optional[str]:
        if not date_text:
            return None
        date_formats = [
            "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y", "%b %d, %Y",
            "%d %B %Y", "%d %b %Y"
        ]
        for fmt in date_formats:
            try:
                parsed_date = datetime.strptime(date_text.strip(), fmt)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        for pattern in get_config(REGEX_CONFIG, "date_patterns"):
            year_month_day = re.search(pattern, date_text)
            if year_month_day:
                groups = year_month_day.groups()
                if len(groups) == 3:
                    year, month, day = groups
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return None
    def _get_mock_metadata(self, certificate_id: str) -> Dict:
        mock = get_config(SELECTORS_CONFIG, "mock_metadata")
        return {
            "certificate_id": certificate_id,
            "name": mock["name"],
            "course": mock["course"],
            "issue_date": mock["issue_date"],
            "platform": get_config(SELECTORS_CONFIG, "platform_names", "udemy"),
            "verification_url": f"https://www.udemy.com/certificate/{certificate_id}"
        }

# --- CERTIFICATE VALIDATION ORCHESTRATOR (from validate_certificate.py) ---
class CertificateValidator:
    """Main certificate validation orchestrator."""
    def __init__(self):
        self.validators: Dict[str, BaseValidator] = {
            'coursera.org': CourseraValidator(),
            'www.coursera.org': CourseraValidator(),
            'credly.com': CredlyValidator(),
            'www.credly.com': CredlyValidator(),
            'udemy.com': UdemyValidator(),
            'www.udemy.com': UdemyValidator(),
            'edx.org': EdXValidator(),
            'www.edx.org': EdXValidator(),
            'linkedin.com': LinkedInValidator(),
            'www.linkedin.com': LinkedInValidator(),
        }
        self._setup_logging()
    def _setup_logging(self):
        log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
        os.makedirs(log_dir, exist_ok=True)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(os.path.join(log_dir, 'validation.log')),
                logging.StreamHandler(sys.stdout)
            ]
        )
        self.logger = logging.getLogger(__name__)
    def _validate_url_format(self, url: str) -> bool:
        try:
            parsed = urlparse(url)
            return (
                parsed.scheme == 'https' and
                parsed.netloc in self.validators and
                len(url) < 2048
            )
        except Exception as e:
            self.logger.error(f"URL validation error: {e}")
            return False
    def _get_validator(self, url: str) -> Optional[BaseValidator]:
        try:
            domain = urlparse(url).netloc
            return self.validators.get(domain)
        except Exception as e:
            self.logger.error(f"Error getting validator: {e}")
            return None
    def validate_certificate(self, url: str, capture_screenshot: bool = False) -> Dict:
        self.logger.info(f"Starting validation for URL: {url}")
        response = {
            "status": "Error",
            "data": {},
            "confidence": 0,
            "screenshot": None,
            "error_message": None
        }
        try:
            if not self._validate_url_format(url):
                response["status"] = "Invalid"
                response["error_message"] = "Invalid URL format or untrusted domain"
                self.logger.warning(f"Invalid URL format: {url}")
                return response
            validator = self._get_validator(url)
            if not validator:
                response["status"] = "Invalid"
                response["error_message"] = "Unsupported platform"
                self.logger.warning(f"Unsupported platform for URL: {url}")
                return response
            if not validator.validate_url_pattern(url):
                response["status"] = "Invalid"
                response["error_message"] = "Invalid URL pattern for platform"
                self.logger.warning(f"Invalid URL pattern: {url}")
                return response
            validation_result = validator.validate_certificate(url, capture_screenshot)
            response.update(validation_result)
            self.logger.info(f"Validation completed for {url}: {response['status']}")
        except Exception as e:
            response["status"] = "Error"
            response["error_message"] = str(e)
            self.logger.error(f"Validation error for {url}: {e}")
        return response

# --- CERTIFICATE VALIDATOR API (from certificate_validator_api.py) ---
class CertificateValidationError(Exception):
    """Custom exception for certificate validation errors."""
    pass

class CertificateValidatorAPI:
    """API wrapper for certificate validation functionality."""
    def __init__(self):
        if CertificateValidator is None:
            raise CertificateValidationError("CertificateValidator not available")
        self.validator = CertificateValidator()
        self.logger = logging.getLogger(__name__)
    def validate_certificates(self, certificate_urls: List[str]) -> Dict[str, Dict]:
        if not certificate_urls:
            return {}
        results = {}
        for url in certificate_urls:
            try:
                self.logger.info(f"Validating certificate: {url}")
                result = self.validator.validate_certificate(url, capture_screenshot=False)
                results[url] = result
            except Exception as e:
                self.logger.error(f"Error validating certificate {url}: {e}")
                results[url] = {
                    "status": "Error",
                    "data": {},
                    "confidence": 0,
                    "screenshot": None,
                    "error_message": str(e)
                }
        return results
    def validate_single_certificate(self, url: str) -> Dict:
        try:
            self.logger.info(f"Validating single certificate: {url}")
            return self.validator.validate_certificate(url, capture_screenshot=False)
        except Exception as e:
            self.logger.error(f"Error validating certificate {url}: {e}")
            return {
                "status": "Error",
                "data": {},
                "confidence": 0,
                "screenshot": None,
                "error_message": str(e)
            }

certificate_validator = None

def get_certificate_validator() -> Optional[CertificateValidatorAPI]:
    global certificate_validator
    if certificate_validator is None:
        try:
            certificate_validator = CertificateValidatorAPI()
        except Exception as e:
            logging.error(f"Failed to initialize certificate validator: {e}")
            return None
    return certificate_validator

def main():
    """CLI interface for certificate validation."""
    import argparse
    parser = argparse.ArgumentParser(description='Certificate Validation System')
    parser.add_argument('url', help='Certificate URL to validate')
    parser.add_argument('--screenshot', action='store_true', help='Capture screenshot of certificate page')
    parser.add_argument('--output', help='Output file for JSON results')
    args = parser.parse_args()
    validator = CertificateValidator()
    result = validator.validate_certificate(args.url, args.screenshot)
    if args.output:
        with open(args.output, 'w') as f:
            json.dump(result, f, indent=2)
        print(f"Results saved to {args.output}")
    else:
        print(json.dumps(result, indent=2))

# CLI entry point for certificate validation---

def clean_text(text: str) -> Optional[str]:
    if not text:
        return None
    return ' '.join(text.split())

def extract_text_from_tags(soup, selectors: List[str], max_length: int = 500, fallback_tags: List[str] = None, keywords: List[str] = None, exclude_keywords: List[str] = None) -> Optional[str]:
    if exclude_keywords is None:
        exclude_keywords = get_config(SELECTORS_CONFIG, "portfolio_exclude_keywords")
    if fallback_tags is None:
        fallback_tags = get_config(SELECTORS_CONFIG, "portfolio_fallback_tags")
    if keywords is None:
        keywords = get_config(SELECTORS_CONFIG, "portfolio_keywords")
    extracted_texts = []
    for selector in selectors:
        tags = soup.select(selector)
        for tag in tags:
            text = clean_text(tag.get_text())
            if text and len(text) <= max_length and (not keywords or any(kw.lower() in text.lower() for kw in keywords)) and not any(ex_kw.lower() in text.lower() for ex_kw in exclude_keywords):
                extracted_texts.append(text)
    if not extracted_texts and fallback_tags:
        for tag_name in fallback_tags:
            tags = soup.find_all(tag_name)
            for tag in tags:
                text = clean_text(tag.get_text())
                if text and len(text) <= max_length and not any(ex_kw.lower() in text.lower() for ex_kw in exclude_keywords) and (not keywords or any(kw.lower() in text.lower() for kw in keywords)):
                    extracted_texts.append(text)
    return " ".join(list(dict.fromkeys(extracted_texts)))[:max_length] if extracted_texts else None

def extract_single_text(soup, selectors: List[str], max_length: int = 500) -> Optional[str]:
    for selector in selectors:
        tag = soup.select_one(selector)
        if tag:
            text = clean_text(tag.get_text())
            if text and len(text) <= max_length:
                return text
    return None

def extract_list_from_tags(soup, selectors: List[str], separator: str = ',') -> List[str]:
    items = []
    exclude_keywords = set(get_config(SELECTORS_CONFIG, "portfolio_list_exclude_keywords"))
    for selector in selectors:
        elements = soup.select(selector)
        for elem in elements:
            text = None
            if elem.name == 'img':
                text = clean_text(elem.get('alt'))
            else:
                text = clean_text(elem.get_text() or elem.get('alt') or elem.get('title'))
            if not text or any(ex_kw.lower() in text.lower() for ex_kw in exclude_keywords):
                continue
            if separator in text:
                items.extend([clean_text(item) for item in text.split(separator) if clean_text(item) and not any(ex_kw.lower() in item.lower() for ex_kw in exclude_keywords)])
            elif len(text.split()) <= 5 and text not in items:
                items.append(text)
    return list(dict.fromkeys([item for item in items if item]))

def extract_link_from_tags(soup, selectors: List[str], base_url: str) -> Optional[str]:
    for selector in selectors:
        tag = soup.select_one(selector)
        if tag and tag.get('href'):
            href = tag['href']
            if href.startswith('mailto:'):
                return href.replace('mailto:', '')
            return urljoin(base_url, href)
    return None

def _find_flexible_tags(soup, keywords, tag_names=None):
    """Find tags where class or id contains any of the keywords (config-driven, no regex)."""
    if tag_names is None:
        tag_names = ["div", "section", "ul", "li", "span", "p"]
    found = []
    for tag in soup.find_all(tag_names):
        classes = ' '.join(tag.get('class', []))
        id_ = tag.get('id', '')
        for kw in keywords:
            if kw.lower() in classes.lower() or kw.lower() in id_.lower():
                found.append(tag)
                break
    return found

def _split_and_clean(text, delimiters=None):
    if not text:
        return []
    if delimiters is None:
        delimiters = get_config(SELECTORS_CONFIG, "portfolio_split_delimiters", default=["\n", ",", "|", "•", "-", "\u2022"])
    # Replace all delimiters with newlines, then split
    for delim in delimiters:
        text = text.replace(delim, "\n")
    items = text.split("\n")
    return [i.strip() for i in items if i and len(i.strip()) > 1]

def parse_portfolio(html_content: str, url: str) -> dict:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # --- Name ---
    name_selectors = get_config(SELECTORS_CONFIG, "portfolio_name_selectors")
    name = extract_single_text(soup, name_selectors)
    if not name:
        og_name = soup.find('meta', property='og:site_name') or soup.find('meta', property='og:title')
        if og_name and og_name.get('content'):
            name = og_name['content'].strip()
        else:
            twitter_name = soup.find('meta', attrs={'name': 'twitter:title'})
            if twitter_name and twitter_name.get('content'):
                name = twitter_name['content'].strip()

    # --- About/Description ---
    about_selectors = get_config(SELECTORS_CONFIG, "portfolio_about_selectors")
    about_max_length = get_config(SELECTORS_CONFIG, "portfolio_about_max_length")
    about_fallback_tags = get_config(SELECTORS_CONFIG, "portfolio_about_fallback_tags")
    about_keywords = get_config(SELECTORS_CONFIG, "portfolio_about_keywords")
    about = extract_text_from_tags(
        soup,
        selectors=about_selectors,
        max_length=about_max_length,
        fallback_tags=about_fallback_tags,
        keywords=about_keywords
    )
    if not about:
        og_desc = soup.find('meta', property='og:description')
        if og_desc and og_desc.get('content'):
            about = og_desc['content'].strip()
        else:
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc and meta_desc.get('content'):
                about = meta_desc['content'].strip()

    # --- Skills ---
    skills = []
    # 1. Try explicit selectors
    skills_selectors = get_config(SELECTORS_CONFIG, "portfolio_skills_selectors")
    skills = extract_list_from_tags(soup, skills_selectors)
    # 2. Try flexible tag search for skills
    if not skills:
        skill_tags = _find_flexible_tags(soup, get_config(SELECTORS_CONFIG, "portfolio_skill_flexible_keywords"))
        for tag in skill_tags:
            skills += _split_and_clean(tag.get_text(" ", strip=True))
    # 3. Try about section
    if not skills and about:
        skills += _split_and_clean(about)
    # 4. Try project descriptions
    if not skills:
        project_tags = _find_flexible_tags(soup, get_config(SELECTORS_CONFIG, "portfolio_project_flexible_keywords"))
        for tag in project_tags:
            skills += _split_and_clean(tag.get_text(" ", strip=True))
    # 5. Filter by known skill keywords
    skill_keywords = set([kw.lower() for kw in get_config(SELECTORS_CONFIG, "skill_keywords")])
    skills = [s for s in set(skills) if any(kw in s.lower() for kw in skill_keywords)]

    # For skills
    skill_min_length = get_config(SELECTORS_CONFIG, "portfolio_skill_min_length", default=2)
    skill_max_length = get_config(SELECTORS_CONFIG, "portfolio_skill_max_length", default=40)
    skills = [
        s for s in set(skills)
        if skill_min_length <= len(s) <= skill_max_length
    ]

    # --- Projects ---
    projects = []
    # 1. Try explicit selectors
    project_selectors = get_config(SELECTORS_CONFIG, "portfolio_project_selectors")
    project_title_selectors = get_config(SELECTORS_CONFIG, "portfolio_project_title_selectors")
    project_desc_selectors = get_config(SELECTORS_CONFIG, "portfolio_project_desc_selectors")
    project_link_selectors = get_config(SELECTORS_CONFIG, "portfolio_project_link_selectors")
    for selector in project_selectors:
        for div in soup.select(selector):
            title = extract_single_text(div, project_title_selectors)
            desc = extract_single_text(div, project_desc_selectors)
            link = extract_link_from_tags(div, project_link_selectors, url)
            if title or desc or link:
                projects.append({
                    "title": title,
                    "description": desc,
                    "link": link
                })
    # 2. Try flexible tag search for projects
    if not projects:
        project_tags = _find_flexible_tags(soup, get_config(SELECTORS_CONFIG, "portfolio_project_flexible_keywords"))
        for tag in project_tags:
            for a in tag.find_all('a', href=True):
                title = a.text.strip() or a.get('title')
                desc = tag.get_text(" ", strip=True)
                projects.append({
                    "title": title,
                    "description": desc,
                    "link": a['href']
                })
    # 3. Fallback: look for links to known project domains
    if not projects:
        fallback_project_domains = get_config(SELECTORS_CONFIG, "portfolio_project_fallback_domains")
        for a in soup.find_all('a', href=True):
            if any(domain in a['href'] for domain in fallback_project_domains):
                title = a.text.strip() or a.get('title')
                desc = None
                parent = a.find_parent(['div', 'li'])
                if parent:
                    desc = parent.get_text(" ", strip=True)
                projects.append({
                    "title": title,
                    "description": desc,
                    "link": a['href']
                })
    # Deduplicate projects by link
    seen_links = set()
    deduped_projects = []
    for proj in projects:
        link = proj.get("link")
        if link and link not in seen_links:
            deduped_projects.append(proj)
            seen_links.add(link)
    projects = deduped_projects

    # --- Education ---
    education = []
    education_selectors = get_config(SELECTORS_CONFIG, "portfolio_education_selectors")
    education = extract_list_from_tags(soup, education_selectors)

    if not education:
        edu_flex_keywords = get_config(SELECTORS_CONFIG, "portfolio_education_flexible_keywords")
        edu_tags = _find_flexible_tags(soup, edu_flex_keywords)
        for tag in edu_tags:
            education += _split_and_clean(tag.get_text(" ", strip=True))

    if not education and about:
        edu_keywords = get_config(SELECTORS_CONFIG, "portfolio_education_keywords")
        education += [line for line in _split_and_clean(about) if any(kw in line.lower() for kw in edu_keywords)]

    # Filter out unwanted items using config
    edu_exclude = [kw.lower() for kw in get_config(SELECTORS_CONFIG, "portfolio_education_exclude_keywords", default=[])]
    edu_min_length = get_config(SELECTORS_CONFIG, "portfolio_education_min_length", default=6)
    edu_max_length = get_config(SELECTORS_CONFIG, "portfolio_education_max_length", default=100)
    education = [
        e for e in dict.fromkeys([e for e in education if edu_min_length <= len(e) <= edu_max_length])
        if not any(ex_kw in e.lower() for ex_kw in edu_exclude)
    ]

    # --- Contact ---
    contact = {}
    contact_selectors = get_config(SELECTORS_CONFIG, "portfolio_contact_selectors")
    for key, sel in contact_selectors.items():
        val = extract_link_from_tags(soup, sel, url)
        if val:
            contact[key] = val

    # --- Experience ---
    experience = []
    experience_keywords = get_config(SELECTORS_CONFIG, "experience_keywords")
    # From about/education
    for line in [about] + education:
        if line and any(word in line.lower() for word in experience_keywords):
            experience.append(line.strip())
    # From projects
    for proj in projects:
        desc = proj.get("description")
        if desc and len(desc.split()) > 3:
            experience.append(desc.strip())
    # Remove duplicates and short lines
    experience = [e for e in dict.fromkeys(experience) if len(e.split()) > 3]

    return {
        "name": name,
        "about": about,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
        "contact": contact
    }

def fetch_portfolio_with_selenium(url: str) -> dict:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from webdriver_manager.chrome import ChromeDriverManager
    import time
    import os
    debug_file = 'selenium_portfolio_debug.html'
    driver = None
    try:
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--window-size=1920,1080')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--no-sandbox')
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)
        driver.get(url)
        # Wait for the page to load (adjust as needed)
        WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
        time.sleep(8)  # Give JS more time to render
        html = driver.page_source
        # Save for debugging
        with open(debug_file, 'w', encoding='utf-8') as f:
            f.write(html)
        result = parse_portfolio(html, url)
        # If nothing extracted, return debug info
        if not any([result.get('name'), result.get('about'), result.get('skills'), result.get('projects'), result.get('education'), result.get('contact')]):
            return {"error": "Portfolio extraction returned no data. See debug HTML.", "debug_html": os.path.abspath(debug_file)}
        return result
    except Exception as e:
        if driver:
            try:
                driver.quit()
            except:
                pass
        return {"error": f"Selenium error: {str(e)}", "debug_html": os.path.abspath(debug_file)}
    finally:
        if driver:
            try:
                driver.quit()
            except:
                pass

def fetch_portfolio(url: str) -> dict:
    # Try Selenium first for JS-rendered content
    try:
        result = fetch_portfolio_with_selenium(url)
        # If Selenium worked and didn't just return an error, use it
        if result and not result.get('error'):
            return result
    except Exception:
        pass
    # Fallback to requests + BeautifulSoup
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        portfolio_info = parse_portfolio(response.text, url)
        return portfolio_info
    except Exception as e:
        return {"error": str(e)}

# --- GITHUB EXTRACTOR LOGIC (from github_extractor.py, functions only) ---

class GitHubAPIError(Exception):
    pass

async def fetch_github_profile(username: str) -> dict:
    user_url = f"https://api.github.com/users/{username}"
    repos_url = f"https://api.github.com/users/{username}/repos"
    events_url = f"https://api.github.com/users/{username}/events/public"
    headers = {
        "Accept": "application/vnd.github.v3+json",
        "User-Agent": "Portfolio-Extractor"
    }

    async with aiohttp.ClientSession() as session:
        # --- Profile Info ---
        async with session.get(user_url, headers=headers) as resp:
            if resp.status != 200:
                return {"error": f"GitHub API returned status {resp.status}"}
            profile = await resp.json()

        # --- Repositories ---
        repos = []
        page = 1
        while True:
            async with session.get(repos_url, headers=headers, params={"page": page, "per_page": 100}) as resp:
                if resp.status != 200:
                    break
                page_repos = await resp.json()
                if not page_repos:
                    break
                for repo in page_repos:
                    repos.append({
                        "name": repo["name"],
                        "description": repo["description"],
                        "language": repo["language"],
                        "stars": repo["stargazers_count"],
                        "forks": repo["forks_count"],
                        "open_issues": repo["open_issues_count"],
                        "watchers": repo["watchers_count"],
                        "size": repo["size"],
                        "created_at": repo["created_at"],
                        "updated_at": repo["updated_at"],
                        "pushed_at": repo["pushed_at"],
                        "url": repo["html_url"],
                        "homepage": repo["homepage"],
                        "topics": repo.get("topics", []),
                        "license": repo["license"]["name"] if repo["license"] else None,
                        "default_branch": repo["default_branch"],
                        "is_fork": repo["fork"],
                        "archived": repo["archived"]
                    })
                if len(page_repos) < 100:
                    break
                page += 1

        # --- Contributions (from events) ---
        async with session.get(events_url, headers=headers, params={"per_page": 100}) as resp:
            contributions = {
                "commits": 0,
                "pull_requests": 0,
                "issues": 0,
                "repositories_contributed_to": set()
            }
            if resp.status == 200:
                events = await resp.json()
                for event in events:
                    if event["type"] == "PushEvent":
                        contributions["commits"] += len(event["payload"]["commits"])
                        contributions["repositories_contributed_to"].add(event["repo"]["name"])
                    elif event["type"] == "PullRequestEvent":
                        contributions["pull_requests"] += 1
                        contributions["repositories_contributed_to"].add(event["repo"]["name"])
                    elif event["type"] == "IssuesEvent":
                        contributions["issues"] += 1
                        contributions["repositories_contributed_to"].add(event["repo"]["name"])
            contributions["repositories_contributed_to"] = list(contributions["repositories_contributed_to"])

        # --- Assemble Output ---
        return {
            "username": profile.get("login"),
            "name": profile.get("name"),
            "bio": profile.get("bio"),
            "location": profile.get("location"),
            "company": profile.get("company"),
            "blog": profile.get("blog"),
            "email": profile.get("email"),
            "twitter_username": profile.get("twitter_username"),
            "public_repos": profile.get("public_repos"),
            "public_gists": profile.get("public_gists"),
            "followers": profile.get("followers"),
            "following": profile.get("following"),
            "created_at": profile.get("created_at"),
            "updated_at": profile.get("updated_at"),
            "avatar_url": profile.get("avatar_url"),
            "hireable": profile.get("hireable"),
            "repositories": repos,
            "contributions": contributions
        }

app = FastAPI(
    title="Your API Title",
    description="Your API Description"
)

# Add CORS middleware to allow frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"message": "Hello World"}

# --- API KEY AUTHORIZATION (single source of truth) ---
API_KEY = "qFvFLN4VeYm3XqxnY0s8p-6isd5FCSF8o5aeuxhyOuw"
API_KEY_NAME = "access_token"
api_key_header = APIKeyHeader(name=API_KEY_NAME, auto_error=False)

def get_api_key(api_key: str = Depends(api_key_header)):
    if api_key == API_KEY:
        return api_key
    else:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or missing API Key",
        )

def extract_elements(text):
    patterns = get_config(REGEX_CONFIG, "extract_elements_patterns")
    github = re.search(patterns["github"], text, re.IGNORECASE)
    instagram = re.search(patterns["instagram"], text, re.IGNORECASE)
    udemy = re.search(patterns["udemy"], text, re.IGNORECASE)
    linkedin = re.search(patterns["linkedin"], text, re.IGNORECASE)
    credly = re.search(patterns["credly"], text, re.IGNORECASE)
    coursera = re.search(patterns["coursera"], text, re.IGNORECASE)

    # Find all URLs in the text
    url_pattern = r"https?://[\w\.-]+(?:/[\w\.-]*)*"
    all_urls = re.findall(url_pattern, text)
    # Remove URLs already classified
    classified = set()
    for match in [github, udemy, linkedin, credly, coursera]:
        if match:
            classified.add(match.group(0))
    portfolio = None
    for url in all_urls:
        if url not in classified:
            portfolio = url
            break

    return {
        "github": github.group(0) if github else None,
        "portfolio": portfolio,
        "instagram": instagram.group(1) if instagram else None,
        "udemy": udemy.group(0) if udemy else None,
        "linkedin": linkedin.group(0) if linkedin else None,
        "credly": credly.group(0) if credly else None,
        "coursera": coursera.group(0) if coursera else None,
    }

@app.post("/upload/")
async def upload_file(
    file: UploadFile = File(...),
    api_key: str = Depends(get_api_key)
):
    # 1. Save and read the uploaded file
    suffix = os.path.splitext(file.filename)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    # 2. Extract text from DOCX
    doc = Document(tmp_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    os.remove(tmp_path)

    # 3. Extract elements (github, portfolio, instagram, etc.)
    elements = extract_elements(text)

    # 3a. Scrape Instagram profile if username found
    instagram_profile = None
    if elements.get("instagram"):
        instagram_profile = parse_instagram(elements["instagram"])

    # 3b. Scrape portfolio site if URL found
    portfolio_data = None
    if elements.get("portfolio"):
        portfolio_data = fetch_portfolio(elements["portfolio"])

    # --- Extract skills and experience if missing ---
    if portfolio_data:
        # Extract skills if empty
        if not portfolio_data.get("skills"):
            skills_set = set()
            about = portfolio_data.get("about", "")
            education = " ".join(portfolio_data.get("education", []))
            projects = " ".join([p.get("description", "") or "" for p in portfolio_data.get("projects", [])])
            text_blob = f"{about} {education} {projects}".lower()
            # Common skill keywords
            skill_keywords = get_config(SELECTORS_CONFIG, "skill_keywords")
            for kw in skill_keywords:
                if kw in text_blob:
                    skills_set.add(kw.title())
            portfolio_data["skills"] = sorted(skills_set)
        # Extract experience if empty
        if not portfolio_data.get("experience"):
            experience = []
            about = portfolio_data.get("about", "")
            education = portfolio_data.get("education", [])
            projects = portfolio_data.get("projects", [])
            experience_keywords = get_config(SELECTORS_CONFIG, "experience_keywords")
            # From about/education
            for line in ([about] + education):
                if any(word in line.lower() for word in experience_keywords):
                    experience.append(line.strip())
            # From projects
            for proj in projects:
                desc = proj.get("description")
                if desc and len(desc.split()) > 3:
                    experience.append(desc.strip())
            # Remove duplicates and short lines
            experience = [e for e in dict.fromkeys(experience) if len(e.split()) > 3]
            portfolio_data["experience"] = experience

    # 3c. Scrape GitHub profile if username or URL found
    github_profile = None
    if elements.get("github"):
        github_url = elements["github"]
        github_username = extract_github_username(github_url)
        if github_username:
            github_profile = await fetch_github_profile(github_username)

    # 4. Collect certificate URLs from both elements and text extraction
    certificate_urls = set(extract_certificate_urls(text))
    for key in ["udemy", "credly", "coursera", "edx", "linkedin"]:
        url = elements.get(key)
        if url:
            certificate_urls.add(url)
    # Filter out any non-URLs (e.g., 'www.')
    certificate_urls = {url for url in certificate_urls if isinstance(url, str) and url.startswith("http")}

    # 5. Validate each certificate URL using the orchestrator
    certificates = {}
    validator = CertificateValidator()
    for url in certificate_urls:
        certificates[url] = validator.validate_certificate(url)

    # 6. Build summary
    total_certificates = len(certificates)
    valid_certificates = sum(1 for c in certificates.values() if c.get("status") == "Valid")
    invalid_certificates = sum(1 for c in certificates.values() if c.get("status") == "Invalid")
    certificate_summary = {
        "total_certificates": total_certificates,
        "valid_certificates": valid_certificates,
        "invalid_certificates": invalid_certificates,
        "validation_rate": int(100 * valid_certificates / total_certificates) if total_certificates else 0
    }

    # 7. Return your response
    return {
        "elements": elements,
        "instagram_profile": instagram_profile,
        "portfolio_data": portfolio_data,
        "github_profile": github_profile,
        "certificates": certificates,
        "certificate_summary": certificate_summary,
    }

async def stream_processing_generator(file_content: bytes, filename: str):
    """Generator function that yields processing updates in real-time"""
    try:
        # 1. Extract text from DOCX
        suffix = os.path.splitext(filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name

        doc = Document(tmp_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        os.remove(tmp_path)

        # Send initial status
        yield f"data: {json.dumps({'type': 'status', 'message': 'Processing document...', 'step': 'extract'})}\n\n"

        # 2. Extract elements
        elements = extract_elements(text)
        yield f"data: {json.dumps({'type': 'elements', 'data': elements})}\n\n"

        # 3a. Scrape Instagram profile if username found
        instagram_profile = None
        if elements.get("instagram"):
            yield f"data: {json.dumps({'type': 'status', 'message': 'Fetching Instagram profile...', 'step': 'instagram'})}\n\n"
            instagram_profile = parse_instagram(elements["instagram"])
            yield f"data: {json.dumps({'type': 'instagram', 'data': instagram_profile})}\n\n"

        # 3b. Scrape portfolio site if URL found
        portfolio_data = None
        if elements.get("portfolio"):
            yield f"data: {json.dumps({'type': 'status', 'message': 'Fetching portfolio data...', 'step': 'portfolio'})}\n\n"
            portfolio_data = fetch_portfolio(elements["portfolio"])
            
            # Extract skills and experience if missing
            if portfolio_data:
                if not portfolio_data.get("skills"):
                    skills_set = set()
                    about = portfolio_data.get("about", "")
                    education = " ".join(portfolio_data.get("education", []))
                    projects = " ".join([p.get("description", "") or "" for p in portfolio_data.get("projects", [])])
                    text_blob = f"{about} {education} {projects}".lower()
                    skill_keywords = get_config(SELECTORS_CONFIG, "skill_keywords")
                    for kw in skill_keywords:
                        if kw in text_blob:
                            skills_set.add(kw.title())
                    portfolio_data["skills"] = sorted(skills_set)
                
                if not portfolio_data.get("experience"):
                    experience = []
                    about = portfolio_data.get("about", "")
                    education = portfolio_data.get("education", [])
                    projects = portfolio_data.get("projects", [])
                    experience_keywords = get_config(SELECTORS_CONFIG, "experience_keywords")
                    for line in ([about] + education):
                        if any(word in line.lower() for word in experience_keywords):
                            experience.append(line.strip())
                    for proj in projects:
                        desc = proj.get("description")
                        if desc and len(desc.split()) > 3:
                            experience.append(desc.strip())
                    experience = [e for e in dict.fromkeys(experience) if len(e.split()) > 3]
                    portfolio_data["experience"] = experience
            
            yield f"data: {json.dumps({'type': 'portfolio', 'data': portfolio_data})}\n\n"

        # 3c. Scrape GitHub profile if username or URL found
        github_profile = None
        if elements.get("github"):
            yield f"data: {json.dumps({'type': 'status', 'message': 'Fetching GitHub profile...', 'step': 'github'})}\n\n"
            github_url = elements["github"]
            github_username = extract_github_username(github_url)
            if github_username:
                github_profile = await fetch_github_profile(github_username)
            yield f"data: {json.dumps({'type': 'github', 'data': github_profile})}\n\n"

        # 4. Collect certificate URLs
        certificate_urls = set(extract_certificate_urls(text))
        for key in ["udemy", "credly", "coursera", "edx", "linkedin"]:
            url = elements.get(key)
            if url:
                certificate_urls.add(url)
        certificate_urls = {url for url in certificate_urls if isinstance(url, str) and url.startswith("http")}

        # 5. Validate each certificate URL
        certificates = {}
        validator = CertificateValidator()
        total_certs = len(certificate_urls)
        
        if total_certs > 0:
            yield f"data: {json.dumps({'type': 'status', 'message': f'Validating {total_certs} certificates...', 'step': 'certificates'})}\n\n"
            
            for i, url in enumerate(certificate_urls, 1):
                yield f"data: {json.dumps({'type': 'status', 'message': f'Validating certificate {i}/{total_certs}...', 'step': 'certificate_progress'})}\n\n"
                certificates[url] = validator.validate_certificate(url)
                yield f"data: {json.dumps({'type': 'certificate_progress', 'url': url, 'data': certificates[url], 'progress': i, 'total': total_certs})}\n\n"

        # 6. Build summary
        total_certificates = len(certificates)
        valid_certificates = sum(1 for c in certificates.values() if c.get("status") == "Valid")
        invalid_certificates = sum(1 for c in certificates.values() if c.get("status") == "Invalid")
        certificate_summary = {
            "total_certificates": total_certificates,
            "valid_certificates": valid_certificates,
            "invalid_certificates": invalid_certificates,
            "validation_rate": int(100 * valid_certificates / total_certificates) if total_certificates else 0
        }

        # 7. Send final results
        final_results = {
            "elements": elements,
            "instagram_profile": instagram_profile,
            "portfolio_data": portfolio_data,
            "github_profile": github_profile,
            "certificates": certificates,
            "certificate_summary": certificate_summary,
        }
        
        yield f"data: {json.dumps({'type': 'complete', 'data': final_results})}\n\n"
        
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

@app.post("/upload/stream/")
async def upload_file_stream(
    file: UploadFile = File(...),
    api_key: str = Depends(get_api_key)
):
    """Streaming endpoint that yields results as they become available"""
    file_content = await file.read()
    return StreamingResponse(
        stream_processing_generator(file_content, file.filename),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "*",
        }
    )

def parse_instagram(username: str) -> dict:
    import requests
    from bs4 import BeautifulSoup
    import re
    from datetime import datetime

    url = f"https://www.instagram.com/{username}/"
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        desc = soup.find("meta", attrs={"name": "description"})
        if desc and desc.get("content"):
            content = desc["content"]
            # Example: '169 followers, 214 following, 1 posts – See Instagram photos and videos from  (@nitish5300)'
            followers = re.search(r"([\d,]+) followers", content, re.IGNORECASE)
            following = re.search(r"([\d,]+) following", content, re.IGNORECASE)
            posts = re.search(r"([\d,]+) posts", content, re.IGNORECASE)
            return {
                "bio": content,
                "followers": int(followers.group(1).replace(',', '')) if followers else None,
                "following": int(following.group(1).replace(',', '')) if following else None,
                "posts_count": int(posts.group(1).replace(',', '')) if posts else None,
                "username": username,
                "timestamp": str(datetime.now())
            }
        else:
            return {"error": "Could not find profile meta description"}
    except Exception as e:
        return {"error": str(e)}

def parse_coursera(url: str) -> dict:
    import requests
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200 and "coursera" in resp.text.lower():
            return {"status": "Valid", "platform": "Coursera", "verification_url": url}
        else:
            return {"status": "Invalid", "platform": "Coursera", "verification_url": url}
    except Exception as e:
        return {"status": "Invalid", "platform": "Coursera", "verification_url": url, "error": str(e)}

def parse_credly(url: str) -> dict:
    import requests
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200 and "badge" in resp.text.lower():
            return {"status": "Valid", "platform": "Credly", "verification_url": url}
        else:
            return {"status": "Invalid", "platform": "Credly", "verification_url": url}
    except Exception as e:
        return {"status": "Invalid", "platform": "Credly", "verification_url": url, "error": str(e)}

def parse_udemy(url: str) -> dict:
    import requests
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200 and "Certificate" in resp.text:
            return {"status": "Valid", "platform": "Udemy", "verification_url": url}
        else:
            return {"status": "Invalid", "platform": "Udemy", "verification_url": url}
    except Exception as e:
        return {"status": "Invalid", "platform": "Udemy", "verification_url": url, "error": str(e)}

def parse_edx(url: str) -> dict:
    import requests
    from bs4 import BeautifulSoup

    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Course title
        course_title = None
        for h1 in soup.find_all("h1"):
            if "certificate" not in h1.text.lower():
                course_title = h1.text.strip()
                break

        # Recipient name
        recipient = None
        for div in soup.find_all("div"):
            if "awarded to" in div.text.lower():
                recipient = div.text.replace("Awarded to", "").strip()
                break

        # Issue date
        issue_date = None
        for div in soup.find_all("div"):
            if "issued" in div.text.lower():
                issue_date = div.text.strip()
                break

        return {
            "course_title": course_title,
            "recipient": recipient,
            "issue_date": issue_date
        }
    except Exception as e:
        return {"error": str(e)}

def extract_github_username(url):
    pattern = get_config(REGEX_CONFIG, "github_username")
    match = re.search(pattern, url)
    return match.group(1) if match else None

def extract_certificate_urls(text):
    urls = []
    for pattern in PATTERNS["certificate_patterns"].values():
        urls += re.findall(pattern, text)
    return urls

def extract_name(soup):
    for selector in get_config(SELECTORS_CONFIG, "name_selectors"):
        tag = soup.select_one(selector)
        if tag and tag.get_text(strip=True):
            return clean_text(tag.get_text())
    # Fallback: largest h1/h2
    candidates = soup.find_all(['h1', 'h2'])
    if candidates:
        return clean_text(max(candidates, key=lambda t: len(t.get_text())).get_text())
    return None

# --- LOAD CONFIGS FOR REGEX AND SELECTORS ---
# CONFIG_DIR = os.path.dirname(os.path.abspath(__file__))
# with open(os.path.join(CONFIG_DIR, 'regex_config.json'), 'r', encoding='utf-8') as f:
#     REGEX_CONFIG = json.load(f)
# with open(os.path.join(CONFIG_DIR, 'selectors_config.json'), 'r', encoding='utf-8') as f:
#     SELECTORS_CONFIG = json.load(f)

# Helper to get config value with error handling
def get_config(config, *keys, default=None):
    try:
        for key in keys:
            config = config[key]
        return config
    except (KeyError, TypeError):
        if default is not None:
            return default
        raise KeyError(f"Missing config value for: {'/'.join(keys)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("complete_file:app", host="0.0.0.0", port=8000, reload=True)
